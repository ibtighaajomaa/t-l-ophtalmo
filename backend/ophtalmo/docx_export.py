import io
import base64
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _parse_html_to_docx(doc: Document, html: str):
    if not html:
        doc.add_paragraph("Contenu non disponible")
        return

    table_pattern = re.compile(r'<table>(.*?)</table>', re.DOTALL)
    parts = table_pattern.split(html)

    for i, part in enumerate(parts):
        if i % 2 == 0:
            _render_text_html(doc, part)
        else:
            _render_table_html(doc, part)


def _render_text_html(doc: Document, html: str):
    lines = re.split(r'<(h[1-6])>(.*?)</\1>|<p>(.*?)</p>|<li>(.*?)</li>|<hr/?>', html)
    i = 0
    while i < len(lines):
        line = lines[i] if i < len(lines) else ""
        if line is None:
            i += 1
            continue
        line = line.strip()
        if not line:
            i += 1
            continue

        if line in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            content = lines[i+1] if i+1 < len(lines) else ""
            level = int(line[1])
            heading = doc.add_heading(content, level=level)
            i += 2
            continue

        stripped = re.sub(r'<[^>]+>', '', line).strip()
        if stripped:
            doc.add_paragraph(stripped)
        i += 1

    ul_match = re.findall(r'<ul>(.*?)</ul>', html, re.DOTALL)
    for ul_content in ul_match:
        items = re.findall(r'<li>(.*?)</li>', ul_content, re.DOTALL)
        for item in items:
            clean = re.sub(r'<[^>]+>', '', item).strip()
            if clean:
                doc.add_paragraph(clean, style='List Bullet')


def _render_table_html(doc: Document, html: str):
    rows = re.findall(r'<tr>(.*?)</tr>', html, re.DOTALL)
    if not rows:
        return
    table_data = []
    for row_html in rows:
        cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row_html, re.DOTALL)
        row = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
        if row:
            table_data.append(row)

    if not table_data:
        return

    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_data in enumerate(table_data):
        for ci, cell_data in enumerate(row_data):
            if ci < len(table.rows[ri].cells):
                cell = table.rows[ri].cells[ci]
                cell.text = cell_data
                if ri == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def _add_header_footer(doc: Document, patient_id: str, report_id: int):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"Tele-Ophtalmo - Rapport Medical - Patient: {patient_id}"
    hp.style.font.size = Pt(8)
    hp.style.font.color.rgb = RGBColor(128, 128, 128)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = f"Rapport genere le {datetime.now():%d/%m/%Y %H:%M} | Document ID: R-{report_id}"
    fp.style.font.size = Pt(8)
    fp.style.font.color.rgb = RGBColor(128, 128, 128)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_report_to_docx(report) -> io.BytesIO:
    doc = Document()

    doc.add_heading("Rapport de Fond d'Œil", level=0)

    doc.add_paragraph(
        f"Patient: {report.patient_id} | Examen: {report.examination_id} | "
        f"Date: {datetime.now():%d/%m/%Y %H:%M}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("—" * 40)

    content = report.final_content or report.doctor_content or report.ai_content
    _parse_html_to_docx(doc, content)

    doc.add_paragraph("—" * 40)

    status_map = {
        "AI_GENERATED": "Genere par IA (en attente de validation)",
        "UNDER_REVIEW": "En cours de revision",
        "SIGNED": "Signe et valide",
    }
    p = doc.add_paragraph()
    p.add_run(f"Statut: {status_map.get(report.status, report.status)}").bold = True

    if report.signed_by:
        p = doc.add_paragraph()
        p.add_run(f"Valide et signe par: {report.signed_by.first_name} {report.signed_by.last_name}").bold = True
    if report.signed_at:
        doc.add_paragraph(f"Date de signature: {report.signed_at:%d/%m/%Y %H:%M}")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run(
        "Ce rapport est genere automatiquement par intelligence artificielle. "
        "Il ne constitue pas un diagnostic medical et doit etre valide par un "
        "ophtalmologue qualifie avant toute decision clinique."
    ).italic = True

    _add_header_footer(doc, report.patient_id, report.pk)

    if report.ai_report_data:
        xai_images = []
        for key in ('gradcam_image', 'clahe_image'):
            val = report.ai_report_data.get(key)
            if val and isinstance(val, str) and len(val) > 100:
                xai_images.append((key, val))
        if xai_images:
            doc.add_page_break()
            doc.add_heading("Annexe - Visualisations XAI", level=1)
            for name, b64data in xai_images:
                doc.add_paragraph(f"Image: {name}")
                try:
                    img_data = base64.b64decode(b64data)
                    img_stream = io.BytesIO(img_data)
                    doc.add_picture(img_stream, width=Inches(3.5))
                except Exception:
                    doc.add_paragraph("(Image non disponible)")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
